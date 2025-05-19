"""
Docx解析器，负责解析Microsoft Word文档并提取文本内容和元数据
"""
import os
from typing import Dict, Any, List, Optional
import docx
from core.utils.logger import get_logger

logger = get_logger(__name__)


class DocxParser:
    """Word文档解析器，提取docx文件的文本内容和元数据"""
    
    def __init__(self):
        """初始化Docx解析器"""
        pass
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析Word文档
        
        Args:
            file_path: Word文档文件路径
            
        Returns:
            Dict: 包含标题、内容和元数据的字典
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return {}
        
        try:
            result = {
                'title': os.path.basename(file_path),
                'content': '',
                'metadata': {}
            }
            
            doc = docx.Document(file_path)
            
            # 提取元数据
            core_properties = doc.core_properties
            if core_properties:
                # 尝试提取常见的文档属性
                metadata = {}
                if core_properties.title:
                    metadata['title'] = core_properties.title
                if core_properties.author:
                    metadata['author'] = core_properties.author
                if core_properties.created:
                    metadata['created'] = str(core_properties.created)
                if core_properties.modified:
                    metadata['modified'] = str(core_properties.modified)
                if core_properties.subject:
                    metadata['subject'] = core_properties.subject
                if core_properties.keywords:
                    metadata['keywords'] = core_properties.keywords
                if core_properties.comments:
                    metadata['comments'] = core_properties.comments
                if core_properties.category:
                    metadata['category'] = core_properties.category
                if core_properties.last_modified_by:
                    metadata['last_modified_by'] = core_properties.last_modified_by
                
                result['metadata'] = metadata
                
                # 如果元数据中有标题，则使用元数据中的标题
                if 'title' in metadata and metadata['title']:
                    result['title'] = metadata['title']
            
            # 提取文档内容
            paragraphs = []
            
            # 提取标题（如果第一段是标题且元数据中没有标题）
            if doc.paragraphs and doc.paragraphs[0].style.name.startswith('Heading') and not result['title']:
                result['title'] = doc.paragraphs[0].text
            
            # 提取所有段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            result['content'] = '\n\n'.join(paragraphs)
            
            # 添加统计信息
            result['metadata']['paragraph_count'] = len(doc.paragraphs)
            result['metadata']['table_count'] = len(doc.tables)
            
            return result
            
        except Exception as e:
            logger.error(f"解析Word文档失败: {file_path}, 错误: {str(e)}")
            return {}
    
    def extract_images(self, file_path: str, output_dir: str = None) -> List[str]:
        """
        从Word文档中提取图片（可选功能）
        
        Args:
            file_path: Word文档文件路径
            output_dir: 图片输出目录，如果为None则不保存
            
        Returns:
            List[str]: 保存的图片文件路径列表
        """
        if not os.path.exists(file_path):
            logger.error(f"文件不存在: {file_path}")
            return []
        
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        image_paths = []
        try:
            doc = docx.Document(file_path)
            
            # 从关系中提取图片
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_ext = rel.target_ref.split(".")[-1]  # 获取图片扩展名
                        
                        if output_dir:
                            filename = f"image_{len(image_paths)+1}.{image_ext}"
                            output_path = os.path.join(output_dir, filename)
                            
                            with open(output_path, "wb") as img_file:
                                img_file.write(image_part.blob)
                            
                            image_paths.append(output_path)
                    except Exception as e:
                        logger.warning(f"提取Word文档图片失败, 错误: {str(e)}")
            
            return image_paths
            
        except Exception as e:
            logger.error(f"从Word文档提取图片失败: {file_path}, 错误: {str(e)}")
            return []